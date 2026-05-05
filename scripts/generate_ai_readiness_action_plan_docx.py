#!/usr/bin/env python3

from __future__ import annotations

import argparse
import importlib.util
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.oxml.ns import nsmap
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DEFAULT_GENERATOR_PATH = Path("scripts/generate_ai_readiness_action_plan.py")
DEFAULT_CONFIG_PATH = Path("configs/audit_config.json")
DEFAULT_STOREFRONT_AUDIT = Path("output/data/shopify_storefront_audit.json")
DEFAULT_ADMIN_AUDIT = Path("output/data/shopify_admin_audit.json")
DEFAULT_TEMPLATE_DOCX = Path("output/doc/ai_readiness_action_plan_2026-03-30.docx")


def load_module(path: Path):
    spec = importlib.util.spec_from_file_location(path.stem, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def apply_table_style(table) -> None:
    for style_name in ("Table Grid", "TableNormal"):
        try:
            table.style = style_name
            return
        except KeyError:
            continue


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    run.bold = True


def add_bullets(document: Document, items: List[str]) -> None:
    for item in items:
        try:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(item)
        except KeyError:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"• {item}")


def paragraph_text_from_xml(node) -> str:
    chunks = [item.text or "" for item in node.findall(".//w:t", namespaces=nsmap)]
    return "".join(chunks).strip()


def reset_document_from_heading(document: Document, heading_text: str) -> None:
    body = document._element.body
    children = list(body)
    start_idx = None
    for idx, child in enumerate(children):
        if child.tag.endswith("}p") and paragraph_text_from_xml(child) == heading_text:
            start_idx = idx
            break
    if start_idx is None:
        return
    for child in children[start_idx:]:
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def set_generated_timestamp(document: Document, generated_at: str) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.lower().startswith("generated:"):
            paragraph.clear()
            run = paragraph.add_run(f"Generated: {generated_at}")
            run.italic = True
            run.font.size = Pt(10)
            return


def make_summary_table(document: Document, rows: List[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    set_cell_text(header[0], "Summary Item", bold=True)
    set_cell_text(header[1], "Value", bold=True)
    set_cell_shading(header[0], "D9EAF7")
    set_cell_shading(header[1], "D9EAF7")
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], value)


def make_priority_table(document: Document, items: List[Dict[str, Any]], findings_by_site: Dict[str, List[Dict[str, Any]]]) -> None:
    table = document.add_table(rows=1, cols=5)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Topic", "Business Impact", "Relevant Brands", "Business Move", "Technical Implementation"]
    for idx, label in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], label, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")

    for item in items:
        business_impact = next(
            (
                finding["why"]
                for site_findings in findings_by_site.values()
                for finding in site_findings
                if finding["key"] == item["key"]
            ),
            "",
        )
        cells = table.add_row().cells
        set_cell_text(cells[0], item["title"])
        set_cell_text(cells[1], business_impact)
        set_cell_text(cells[2], ", ".join(item["sites"]))
        set_cell_text(cells[3], item["actions"][0] if item["actions"] else item["title"])
        technical = [f"Lead: {item['team']}"] + [f"- {action}" for action in item["actions"]]
        if item.get("optional"):
            technical.insert(1, "- Optional after core fixes are stable.")
        set_cell_text(cells[4], "\n".join(technical))


def make_site_table(document: Document, items: List[Dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=6)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Priority", "Topic", "Business Impact", "Current Position", "Business Move", "Technical Implementation"]
    for idx, label in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], label, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8F1E8")

    for item in items:
        cells = table.add_row().cells
        set_cell_text(cells[0], item["priority"])
        set_cell_text(cells[1], item["title"])
        set_cell_text(cells[2], item["why"])
        set_cell_text(cells[3], item["evidence"])
        set_cell_text(cells[4], item["recommendation"])
        technical_lines = [f"Owner: {item['team']}", f"Area: {item['area']}"] + [f"- {action}" for action in item["actions"]]
        if item.get("optional"):
            technical_lines.insert(2, "- Optional after higher-priority fixes.")
        set_cell_text(cells[5], "\n".join(technical_lines))


def make_effort_table(document: Document, rows: List[tuple[str, str, str]]) -> None:
    table = document.add_table(rows=1, cols=4)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "Metric", "Estimate", "Estimate Basis"]
    for idx, label in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], label, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")

    for row_idx, (metric, estimate, basis) in enumerate(rows, start=1):
        cells = table.add_row().cells
        set_cell_text(cells[0], str(row_idx))
        set_cell_text(cells[1], metric)
        set_cell_text(cells[2], estimate)
        set_cell_text(cells[3], basis)


def make_consolidated_task_table(document: Document, actions: List[Dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=6)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "Action", "Domains", "Priority", "Estimate", "Expected Outcome"]
    for idx, label in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], label, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")

    for row_idx, item in enumerate(actions, start=1):
        cells = table.add_row().cells
        set_cell_text(cells[0], str(row_idx))
        set_cell_text(cells[1], item["action"])
        set_cell_text(cells[2], item["domains"])
        set_cell_text(cells[3], item["priority"])
        set_cell_text(cells[4], f'{item["estimate"]} ({item["estimate_basis"]})')
        set_cell_text(cells[5], item["expected_outcome"])


def make_timeline_table(document: Document) -> None:
    rows = [
        (
            "Phase 1",
            "Discovery, backlog shaping, implementation planning, and the first pass of homepage metadata, homepage entity/schema, and core schema/page-alignment fixes.",
            "The program starts with an agreed roadmap and the most important brand and site signals begin to stabilize.",
        ),
        (
            "Phase 2",
            "PDP schema expansion, product metadata cleanup, breadcrumb improvements, stronger trust, help, shipping, and policy discoverability, plus the first biweekly visibility re-test after structured-data changes.",
            "Product pages become cleaner machine-readable assets and easier for AI systems to interpret accurately, with early visibility feedback starting to validate the direction of change.",
        ),
        (
            "Phase 3",
            "Richer PDP content rollout, apparel attribute depth, related-product logic, follow-up-answer content, secondary feed or structured-output improvements, and continued biweekly visibility re-tests after schema/content releases.",
            "The storefronts become more citation-worthy and more capable of answering real shopper follow-up questions, while visibility tracking shows whether those updates are converting into stronger recall and citations.",
        ),
        (
            "Phase 4",
            "Brand-specific gap closure for the lower-scoring stores, plus refinement of remaining schema, metadata, and content gaps that were deprioritized earlier, with biweekly visibility re-tests continuing after major JSON-LD updates and custom MCP discovery or architecture work starting.",
            "Lower-scoring stores close their larger readiness gaps while stronger stores get refinement rather than rework, and the custom MCP workstream starts inside the same overall program.",
        ),
        (
            "Phase 5",
            "Biweekly visibility re-testing, storefront re-audit, implementation QA, stakeholder review of what improved, and the main custom MCP design and build phase.",
            "The team can verify whether citation coverage, brand recall, and storefront readiness materially improved, while the custom MCP moves from requirements into active implementation.",
        ),
        (
            "Phase 6",
            "Custom MCP build completion, QA, pilot rollout, final optimisation cycle, content polish, documentation, repeatable QA handoff, and prioritisation of any follow-on items.",
            "The engagement closes with storefront-readiness improvements and a custom MCP workstream delivered inside the same delivery program.",
        ),
    ]
    table = document.add_table(rows=1, cols=3)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Phase", "Focus", "Expected Outcome"]
    for idx, label in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], label, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8F1E8")
    for timeline, focus, outcome in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], timeline)
        set_cell_text(cells[1], focus)
        set_cell_text(cells[2], outcome)


def make_custom_mcp_table(document: Document, relevant_brands: str) -> None:
    table = document.add_table(rows=1, cols=4)
    apply_table_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Business View", "Lead", "Technical Implementation", "Effort"]
    for idx, label in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], label, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")

    cells = table.add_row().cells
    set_cell_text(
        cells[0],
        f"Why this matters: Shopify's native Storefront MCP is useful for baseline commerce tasks, but Shopify does not let you edit that native MCP to add brand-specific advisory logic.\nRelevant brands: {relevant_brands}\nRecommended business move: If AI-assisted shopping is a priority, plan a separate custom MCP for recommendation reasons, comparison logic, sizing or fit advice, brand facts, and market-specific guidance.",
    )
    set_cell_text(cells[1], "Product + Dev")
    set_cell_text(
        cells[2],
        "Keep Shopify's native MCP for catalog search, cart operations, product lookup, and policy answers.\nUse a separate custom MCP only for logic Shopify's native MCP does not cover well enough, such as recommendation explanations, product comparison, sizing or fit interpretation, and brand guidance.\nDo not rebuild standard catalog or cart tools inside the custom MCP if Shopify's native MCP already handles them.",
    )
    set_cell_text(cells[3], "50-55 mandays approx (Dev 40-44 + QA 10-11)")


def build_document(
    output_path: Path,
    findings_by_site: Dict[str, List[Dict[str, Any]]],
    cross_site: List[Dict[str, Any]],
    consolidated_actions: List[Dict[str, Any]],
    effort_summary: List[tuple[str, str, str]],
    generated_at: str,
    visibility_report_path: str | None,
    include_custom_mcp_note: bool,
    template_docx: Path | None = None,
) -> None:
    site_labels: List[str] = []
    for domain, items in findings_by_site.items():
        label = items[0]["site_label"] if items else domain
        if label not in site_labels:
            site_labels.append(label)
    relevant_brands = ", ".join(site_labels) if site_labels else "the audited brands"

    if template_docx and template_docx.exists():
        document = Document(template_docx)
        set_generated_timestamp(document, generated_at)
        reset_document_from_heading(document, "Executive Summary")
    else:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("AI Readiness Action Plan")
        run.bold = True
        run.font.size = Pt(20)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"Generated: {generated_at}")
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(10)

    add_heading(document, "Executive Summary", 1)
    p0_count = sum(1 for item in cross_site if item["priority"] == "P0")
    p1_count = sum(1 for item in cross_site if item["priority"] == "P1")
    workdays_per_week = 5
    project_weeks = 36
    project_capacity = workdays_per_week * project_weeks
    effort_map = {label: value for label, value, _basis in effort_summary}
    summary_rows = [
        ("Sites reviewed", str(len(findings_by_site))),
        ("Immediate priorities that need action now", str(p0_count)),
        ("Secondary priorities to tackle after the foundations are fixed", str(p1_count)),
        ("Planned engagement model", f"6 months at {workdays_per_week} working days per week"),
        ("Indicative project capacity over that period", f"{project_capacity} working days"),
        ("Estimated core storefront-readiness effort", effort_map.get("Core storefront-readiness effort", "-")),
    ]
    if include_custom_mcp_note:
        summary_rows.append(("Estimated custom MCP effort", effort_map.get("Custom MCP effort", "-")))
        summary_rows.append(("Estimated total effort including custom MCP", effort_map.get("Total effort including custom MCP", "-")))
    summary_rows.append(("Estimate basis", "Includes development plus QA or re-test time. Excludes stakeholder waiting time, approval delays, and long content-production cycles."))
    if visibility_report_path:
        summary_rows.append(("Visibility input workbook", visibility_report_path))
    make_summary_table(document, summary_rows)

    add_heading(document, "Estimated Delivery Effort", 1)
    make_effort_table(document, effort_summary)

    add_heading(document, "Consolidated Task List", 1)
    intro = document.add_paragraph()
    intro.add_run("All task estimates below include implementation plus QA or re-test time.").italic = True
    make_consolidated_task_table(document, consolidated_actions)

    add_heading(document, "Cross-Site Priorities", 1)
    for priority, heading in [("P0", "Immediate Priorities"), ("P1", "Next Priorities"), ("P2", "Optional Opportunities")]:
        items = [item for item in cross_site if item["priority"] == priority]
        if not items:
            continue
        add_heading(document, heading, 1)
        make_priority_table(document, items, findings_by_site)

    if include_custom_mcp_note:
        add_heading(document, "Strategic Build Included In This Plan: Custom MCP", 1)
        make_custom_mcp_table(document, relevant_brands)

    add_heading(document, "Site-by-Site Action Plan", 1)
    for domain, items in findings_by_site.items():
        site_label = items[0]["site_label"] if items else domain
        add_heading(document, site_label, 2)
        make_site_table(document, items)

    add_heading(document, "Conclusion", 1)
    conclusion_items = [
        f"Most of the important AI-readiness foundation work applies across all {len(site_labels)} brands, not just one site.",
        "The biggest shared themes are AI product citation coverage, homepage schema depth, homepage title/meta/H1 consistency, deeper product and collection schema, schema-to-page alignment, trust/help discoverability, and AI brand visibility.",
        f"The implementation scope is not identical across the {len(site_labels)} stores, so the rollout should use one shared foundation roadmap plus brand-specific fixes.",
        "The per-site action tables should drive sequencing: apply shared P0 fixes first, then close larger store-specific PDP and schema gaps.",
    ]
    if any("aari" in label.lower() for label in site_labels):
        conclusion_items.append(
            "Aari Clothing is included as an incremental workstream in the same roadmap and should be tracked with the same QA and re-test cadence."
        )
    add_bullets(document, conclusion_items)

    add_heading(document, "Delivery Phases", 1)
    p = document.add_paragraph()
    p.add_run("Visibility re-testing should run every 2 weeks after JSON-LD or other structured-data updates so the team can measure whether schema changes are improving brand recall and product citation outcomes.")
    make_timeline_table(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--storefront-audit", default=str(DEFAULT_STOREFRONT_AUDIT))
    parser.add_argument("--admin-audit", default=str(DEFAULT_ADMIN_AUDIT))
    parser.add_argument("--visibility-report", default="")
    parser.add_argument("--config", default=str(DEFAULT_CONFIG_PATH))
    parser.add_argument("--template-docx", default=str(DEFAULT_TEMPLATE_DOCX))
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    generator = load_module(DEFAULT_GENERATOR_PATH)
    storefront_path = Path(args.storefront_audit)
    admin_path = Path(args.admin_audit)
    visibility_path = Path(args.visibility_report) if args.visibility_report else None
    config_path = Path(args.config)
    template_docx = Path(args.template_docx) if args.template_docx else None

    storefront_payload = generator.load_json(storefront_path)
    admin_payload = generator.load_json(admin_path) if admin_path.exists() else None
    visibility_payload = generator.load_visibility_workbook(visibility_path) if visibility_path and visibility_path.exists() else None
    config = generator.load_config(config_path)
    labels = generator.label_map(config, storefront_payload.get("results", []))
    storefront_module = generator.load_existing_module(Path("scripts/generate_shopify_audit_workbook.py"))

    findings_by_site = generator.build_site_findings(storefront_payload, admin_payload, storefront_module, labels)
    visibility_findings = generator.build_visibility_findings(visibility_payload, config, labels)
    for domain, items in visibility_findings.items():
        findings_by_site[domain] = sorted(findings_by_site.get(domain, []) + items, key=generator.severity_key)
    cross_site = generator.build_cross_site_summary(findings_by_site)
    site_results = storefront_payload.get("results", [])
    site_result_maps = [
        {check["key"]: storefront_module.eval_check(check["key"], storefront_module.build_metrics(site)) for check in storefront_module.DETAIL_CHECKS}
        for site in site_results
    ]
    consolidated_actions = storefront_module.build_consolidated_actions(site_results, site_result_maps)
    effort_summary = storefront_module.consolidated_action_effort_summary(consolidated_actions)
    include_custom_mcp_note = True
    generated_at = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"

    build_document(
        Path(args.output),
        findings_by_site,
        cross_site,
        consolidated_actions,
        effort_summary,
        generated_at,
        str(visibility_path) if visibility_path and visibility_path.exists() else None,
        include_custom_mcp_note,
        template_docx,
    )


if __name__ == "__main__":
    main()
